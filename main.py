import streamlit as st
import google.generativeai as genai
import os
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import pandas as pd
import docx2txt
import fitz  

# ---------------- CONFIG ----------------
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
model = genai.GenerativeModel("gemini-1.5-flash")

# ---------------- STATE INIT ----------------
if "topics" not in st.session_state:
    st.session_state.topics = {"Default": []}
if "current_topic" not in st.session_state:
    st.session_state.current_topic = "Default"
if "language" not in st.session_state:
    st.session_state.language = "Indonesia"
if "analysis_result" not in st.session_state:
    st.session_state.analysis_result = None

# ---------------- SIDEBAR ----------------
st.sidebar.title("⚙️ Pengaturan")

lang_choice = st.sidebar.radio("Pilih Bahasa", ["Indonesia", "English"])
st.session_state.language = lang_choice

st.sidebar.subheader("🗂 Manajemen Topik")
topic_names = list(st.session_state.topics.keys())
selected_topic = st.sidebar.selectbox(
    "Pilih Topik", topic_names, index=topic_names.index(st.session_state.current_topic)
)
st.session_state.current_topic = selected_topic

new_topic = st.sidebar.text_input("Buat Topik Baru")
if st.sidebar.button("➕ Tambah Topik") and new_topic:
    if new_topic not in st.session_state.topics:
        st.session_state.topics[new_topic] = []
        st.session_state.current_topic = new_topic

# ---------------- MAIN ----------------
st.title("Chatbot Industri Value Chain (IVCA)")

# ---- Upload File ----
uploaded_file = st.file_uploader("📂 Upload laporan perusahaan (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
file_content = None

if uploaded_file:
    if uploaded_file.type == "application/pdf":
        pdf = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text = ""
        for page in pdf:
            text += page.get_text()
        file_content = text

    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = docx2txt.process(uploaded_file)
        file_content = text

    elif uploaded_file.type == "text/plain":
        text = uploaded_file.read().decode("utf-8")
        file_content = text

    if file_content:
        st.success("✅ File berhasil dibaca, siap dianalisis.")

        if st.button("🔎 Analisis Value Chain dari Dokumen"):
            instruction = f"""
            Dari dokumen berikut, identifikasi terlebih dahulu objek perusahaan yang dianalisis
            (nama perusahaan, bidang usaha, atau gambaran umum berdasarkan teks).
            Setelah itu lanjutkan dengan analisis Value Chain, Masalah, Solusi, dan Proposal.

            Dokumen: {file_content[:5000]}

            Jawaban harus berformat (bahasa {st.session_state.language}):

            ### Objek Perusahaan
            - Nama/Deskripsi perusahaan
            - Bidang industri
            - Gambaran umum operasional

            ### Value Chain
            #### Aktivitas Utama
            - ...
            - ...
            #### Aktivitas Pendukung
            - ...
            - ...

            ### Masalah
            - ...

            ### Solusi
            - ...

            ### Proposal
            1. Latar Belakang
            2. Identifikasi Masalah
            3. Solusi yang Diusulkan
            4. Rekomendasi
            """
            response = model.generate_content(instruction)
            answer = response.text
            st.session_state.analysis_result = answer
            st.session_state.topics[st.session_state.current_topic].append(("assistant", answer))
            st.markdown(answer)

# ---- Chat Mode ----
for role, msg in st.session_state.topics[st.session_state.current_topic]:
    if role == "user":
        st.chat_message("user").write(msg)
    else:
        st.chat_message("assistant").markdown(msg)

prompt = st.chat_input("Ketik pertanyaan atau perintah...")
if prompt:
    st.session_state.topics[st.session_state.current_topic].append(("user", prompt))
    st.chat_message("user").write(prompt)

    if "value chain" in prompt.lower() or "list of company pain problems and possible solutions" in prompt.lower():
        instruction = f"""
        Pertama-tama, jelaskan objek perusahaan yang dianalisis (nama, bidang, gambaran umum).
        Setelah itu lanjutkan analisis sesuai prompt berikut: {prompt}.

        Jawaban harus berformat (bahasa {st.session_state.language}):

        ### Objek Perusahaan
        - Nama/Deskripsi perusahaan
        - Bidang industri
        - Gambaran umum operasional

        ### Value Chain
        #### Aktivitas Utama
        - ...
        #### Aktivitas Pendukung
        - ...

        ### Masalah
        - ...
        ### Solusi
        - ...
        ### Proposal
        1. Latar Belakang
        2. Identifikasi Masalah
        3. Solusi yang Diusulkan
        4. Rekomendasi
        """
        response = model.generate_content(instruction)
        answer = response.text
        st.session_state.analysis_result = answer
    else:
        if st.session_state.language == "Indonesia":
            response = model.generate_content(f"Jawablah dalam bahasa Indonesia: {prompt}")
        else:
            response = model.generate_content(f"Answer in English: {prompt}")
        answer = response.text

    st.session_state.topics[st.session_state.current_topic].append(("assistant", answer))
    st.chat_message("assistant").markdown(answer)

# ---------------- EXPORT ----------------
st.subheader("📑 Ekspor Proposal / Hasil Analisis")

def export_to_docx(topic_name, text):
    doc = Document()
    doc.add_heading(f"Hasil Analisis - {topic_name}", level=1)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_pdf(topic_name, text):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    content = [Paragraph(f"<b>Hasil Analisis - {topic_name}</b>", styles["Heading1"])]
    content.append(Paragraph(text.replace("\n", "<br/>"), styles["Normal"]))
    doc.build(content)
    buffer.seek(0)
    return buffer

def export_to_excel(topic_name, text):
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        data = []
        for role, msg in st.session_state.topics[st.session_state.current_topic]:
            data.append({"Role": role, "Message": msg})
        df_chat = pd.DataFrame(data)
        df_chat.to_excel(writer, sheet_name="Chat History", index=False)

        if text:
            lines = text.splitlines()
            primary, support, problems, solutions = [], [], [], []
            section = None
            for line in lines:
                if "Aktivitas Utama" in line:
                    section = "primary"
                elif "Aktivitas Pendukung" in line:
                    section = "support"
                elif "Masalah" in line:
                    section = "problem"
                elif "Solusi" in line:
                    section = "solution"
                elif line.strip().startswith("-") or line.strip().startswith("1."):
                    content = line.strip("-1234567890. ").strip()
                    if section == "primary": primary.append(content)
                    elif section == "support": support.append(content)
                    elif section == "problem": problems.append(content)
                    elif section == "solution": solutions.append(content)

            df_vc = pd.DataFrame({
                "Aktivitas Utama": primary + [""]*(max(len(support), len(problems), len(solutions)) - len(primary)),
                "Aktivitas Pendukung": support + [""]*(max(len(primary), len(problems), len(solutions)) - len(support)),
                "Masalah": problems + [""]*(max(len(primary), len(support), len(solutions)) - len(problems)),
                "Solusi": solutions + [""]*(max(len(primary), len(support), len(problems)) - len(solutions)),
            })
            df_vc.to_excel(writer, sheet_name="Value Chain", index=False)

        if text:
            df_proposal = pd.DataFrame({"Proposal": [text]})
            df_proposal.to_excel(writer, sheet_name="Proposal", index=False)

    buffer.seek(0)
    return buffer

if st.session_state.analysis_result:
    st.success("✅ Analisis tersedia, ekspor sekarang!")

    col1, col2, col3 = st.columns(3)
    with col1:
        docx_data = export_to_docx(st.session_state.current_topic, st.session_state.analysis_result)
        st.download_button("⬇️ Word", docx_data, file_name="proposal.docx")
    with col2:
        pdf_data = export_to_pdf(st.session_state.current_topic, st.session_state.analysis_result)
        st.download_button("⬇️ PDF", pdf_data, file_name="proposal.pdf")
    with col3:
        excel_data = export_to_excel(st.session_state.current_topic, st.session_state.analysis_result)
        st.download_button("⬇️ Excel", excel_data, file_name="proposal.xlsx")
